"""
COF (Certificate of Facts) generation — upload-once server-workbook model.
The team uploads the tracking workbook through the browser; it's stored on
the server as the single active workbook. Each generation mutates that
stored file in place under a lock, so COF numbering can't collide.
"""
import contextlib
import datetime
import io
import os
import re
import time

from django.conf import settings

from num2words import num2words
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


CONSIGNOR_NAME  = "PIAGGIO VEHICLES PRIVATE LIMITED"
CONSIGNOR_ADDR  = "F-3 KATHFAL ROAD, MIDC AREA, BARAMATI"
CONSIGNOR_STATE = "MAHARASHTRA"

REMARK_OPTIONS      = ["Short", "Damaged", "Lost", "Missing", "Partial Delivery"]
BIZ_OPTIONS         = ["CV", "2W"]
TRANSPORTER_OPTIONS = ["Delhivery", "DTDC"]

DATA_SHEET = "Data Sheet"


class WorkbookInvalid(Exception):
    """The workbook isn't usable (missing sheets / not an xlsx)."""


class WorkbookInUse(Exception):
    """The workbook file is locked by the OS (e.g. open in Excel on the server)."""


class COFLockTimeout(Exception):
    """Couldn't acquire the generation lock (another COF is in progress)."""


class AssetMissing(Exception):
    """The letterhead template is missing on the server."""


# ── Helpers ──────────────────────────────────────────────────────────
def amount_to_words(amount):
    try:
        amount = float(amount)
        rupees = int(amount)
        paise  = round((amount - rupees) * 100)
        words  = num2words(rupees, lang="en_IN").replace("-", " ").title()
        if paise:
            pw = num2words(paise, lang="en_IN").replace("-", " ").title()
            return f"{words} Rupees And {pw} Paise Only"
        return f"{words} Rupees Only"
    except Exception:
        return str(amount)


def parse_date(s):
    if not s:
        return None
    if isinstance(s, datetime.datetime):
        return s
    if isinstance(s, datetime.date):
        return datetime.datetime(s.year, s.month, s.day)
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%d/%m/%y", "%d-%m-%y", "%Y-%m-%d"):
        try:
            return datetime.datetime.strptime(str(s).strip(), fmt)
        except ValueError:
            pass
    return None


def build_cert_text(remark, loss, words):
    short = "short" in str(remark).lower()
    kind = "short" if short else "damage"
    connector = "consignor / consignee" if short else "consignor/consignee"
    return (
        f"This is to certify that the material {kind} part Value is Rs {loss}"
        f"( {words} ) that the same was insured against transit risk by the "
        f"{connector} with their underwriters. All due care and precaution were "
        f"undertaken for the safe transit. The above loss is not attributable to any "
        f"negligence whatsoever direct, indirect or contributory either on our part or on "
        f"the part of our employees or agents. The Certificate of Facts is issued to enable "
        f"the customer to claim the losses if any, from their underwriters through whom the "
        f"consignment was insured."
    )


# ── Validation & numbering (work off a file path) ────────────────────
def validate_workbook(path):
    """Raise WorkbookInvalid if the file isn't a usable tracking workbook."""
    try:
        wb = load_workbook(path, read_only=True)
    except Exception as e:
        raise WorkbookInvalid(f"Couldn't open that file as an Excel workbook ({e}).")
    try:
        names = wb.sheetnames
    finally:
        wb.close()
    if DATA_SHEET not in names:
        raise WorkbookInvalid(
            "This doesn't look like the COF tracking workbook (its data log sheet is missing).")
    if settings.COF_TEMPLATE_SHEET not in names:
        raise WorkbookInvalid(
            "This doesn't look like the COF tracking workbook (its blank COF template sheet is missing).")


def get_next_cof_info(path):
    wb = load_workbook(path, read_only=True)
    ws = wb[DATA_SHEET]
    last_serial, last_efe = 0, 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] and isinstance(row[0], (int, float)):
            last_serial = int(row[0])
        if len(row) > 14 and row[14]:
            val = str(row[14]).strip()
            if val.upper().startswith("EFE-"):
                try:
                    n = int(val.split("-")[1])
                    last_efe = max(last_efe, n)
                except (ValueError, IndexError):
                    pass
    wb.close()
    nxt = last_efe + 1
    return {
        "serial": last_serial + 1,
        "cof_number": f"EFE-COF-{nxt:04d}",
        "claim_number": nxt,
        "optlog": f"EFE-{nxt:03d}",
    }


# ── Excel mutation ───────────────────────────────────────────────────
def create_cof_sheet(wb, d):
    template_sheet = settings.COF_TEMPLATE_SHEET
    base = d["consignee_name"][:31].strip()
    name, ctr = base, 2
    while name in wb.sheetnames:
        name = f"{base[:28]}{ctr}"
        ctr += 1
    ws = wb.copy_worksheet(wb[template_sheet])
    ws.title = name
    today = datetime.date.today()
    ws["A2"].value = "COF NUMBER";   ws["B2"].value = d["cof_number"]
    ws["A3"].value = "CLAIM NUMBER"; ws["B3"].value = d["claim_number"]
    ws["A4"].value = "DATED";        ws["B4"].value = today
    ws["B4"].number_format = "DD-MMM-YYYY"
    ws["A7"].value = CONSIGNOR_NAME;  ws["C7"].value = d["consignee_name"]
    ws["A8"].value = CONSIGNOR_ADDR;  ws["C8"].value = d["consignee_address"]
    ws["A9"].value = CONSIGNOR_STATE; ws["C9"].value = d["consignee_state"]
    ws["A12"].value = d["invoice_numbers"]
    bd = parse_date(d["invoice_date"])
    ws["B12"].value = bd if bd else d["invoice_date"]
    if bd:
        ws["B12"].number_format = "DD-MMM-YYYY"
    ws["C12"].value = d["num_packages"]; ws["D12"].value = d["weight"]
    ws["E12"].value = d["weight"];       ws["F12"].value = "Carton"
    ws["G12"].value = "Spare Parts"
    ws["A15"].value = "Baramati";    ws["B15"].value = d["destination_city"]
    dd = parse_date(d["delivery_date"])
    ws["C15"].value = dd if dd else d["delivery_date"]
    if dd:
        ws["C15"].number_format = "DD-MMM-YYYY"
    ws["D15"].value = d["lr_number"]
    loss = d["loss_amount"]; words = amount_to_words(loss)
    ws["A17"].value = build_cert_text(d["remark"], loss, words)
    ws["A17"].alignment = Alignment(wrap_text=True, horizontal="left")
    ws["A18"].value = "This is Electronically Generated Certificate, Signature Not Required."
    ws["A18"].font = Font(bold=True, italic=True, size=12, name="Calibri")
    return name


def append_data_sheet(wb, d, serial):
    ws = wb[DATA_SHEET]
    next_row = 2
    for row in ws.iter_rows(min_row=2):
        if row[0].value is not None:
            next_row = row[0].row + 1
    today = datetime.date.today()
    pd = parse_date(d["pickup_date"])
    ws.cell(next_row, 1).value = serial
    ws.cell(next_row, 2).value = d["lr_number"]
    ws.cell(next_row, 3).value = pd if pd else d["pickup_date"]
    if pd:
        ws.cell(next_row, 3).number_format = "DD-MMM-YYYY"
    ws.cell(next_row, 4).value = d["invoice_numbers"]
    ws.cell(next_row, 5).value = d["remark"]
    ws.cell(next_row, 6).value = d["dealer_name"]
    ws.cell(next_row, 7).value = d["state"]
    ws.cell(next_row, 8).value = d["loss_amount"]
    ws.cell(next_row, 9).value = ""
    ws.cell(next_row, 10).value = ""
    ws.cell(next_row, 11).value = ""
    ws.cell(next_row, 12).value = d["status_delhivery"]
    ws.cell(next_row, 13).value = ""
    ws.cell(next_row, 14).value = today
    ws.cell(next_row, 14).number_format = "DD-MMM-YYYY"
    ws.cell(next_row, 15).value = d["optlog"]
    ws.cell(next_row, 16).value = d["ref_delhivery"]


# ── Word generation (in memory) ──────────────────────────────────────
def build_word_doc(d):
    doc = Document(str(settings.COF_LETTERHEAD_PATH))
    today = datetime.date.today().strftime("%d-%b-%Y")
    loss = d["loss_amount"]; words = amount_to_words(loss)

    def cs(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]; p.alignment = align
        r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CERTIFICATE OF FACTS"); r.bold = True; r.font.size = Pt(14)
    doc.add_paragraph()

    t = doc.add_table(rows=3, cols=2); t.style = "Table Grid"
    for i, (l, v) in enumerate([("COF NUMBER", d["cof_number"]),
                                ("CLAIM NUMBER", str(d["claim_number"])), ("DATED", today)]):
        cs(t.cell(i, 0), l, bold=True); cs(t.cell(i, 1), v)
    doc.add_paragraph()

    t2 = doc.add_table(rows=4, cols=2); t2.style = "Table Grid"
    for i, (l, r_) in enumerate([("CONSIGNOR", "CONSIGNEE"),
                                 (CONSIGNOR_NAME, d["consignee_name"]),
                                 (CONSIGNOR_ADDR, d["consignee_address"]),
                                 (CONSIGNOR_STATE, d["consignee_state"])]):
        cs(t2.cell(i, 0), l, bold=(i == 0)); cs(t2.cell(i, 1), r_, bold=(i == 0))
    doc.add_paragraph()

    inv = doc.add_table(rows=2, cols=7); inv.style = "Table Grid"
    bd_dt = parse_date(d["invoice_date"])
    bd_d = bd_dt.strftime("%d-%b-%Y") if bd_dt else d["invoice_date"]
    for i, h in enumerate(["INVOICE NO", "DATED", "NOS OF BOX", "BOOKING WEIGHT",
                           "DELIVERY WEIGHT", "PACKING", "CONTAINS"]):
        cs(inv.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, v in enumerate([d["invoice_numbers"], bd_d, str(d["num_packages"]),
                           str(d["weight"]), str(d["weight"]), "Carton", "Spare Parts"]):
        cs(inv.cell(1, i), v, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    rt = doc.add_table(rows=2, cols=4); rt.style = "Table Grid"
    dd_dt = parse_date(d["delivery_date"])
    dd_d = dd_dt.strftime("%d-%b-%Y") if dd_dt else d["delivery_date"]
    for i, h in enumerate(["FROM", "TO", "DATE OF DELIVERY", "LR NO"]):
        cs(rt.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, v in enumerate(["Baramati", d["destination_city"], dd_d, str(d["lr_number"])]):
        cs(rt.cell(1, i), v, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(build_cert_text(d["remark"], loss, words)); r.font.size = Pt(11)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This is Electronically Generated Certificate, Signature Not Required.")
    r.bold = True; r.italic = True; r.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Lock around read-number → mutate → save ──────────────────────────
@contextlib.contextmanager
def workbook_lock(timeout=20, poll=0.3, stale=120):
    lock_path = str(settings.COF_LOCK_PATH)
    os.makedirs(os.path.dirname(lock_path), exist_ok=True)
    deadline = time.monotonic() + timeout
    fd = None
    while True:
        try:
            fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_RDWR)
            break
        except FileExistsError:
            try:
                if time.time() - os.path.getmtime(lock_path) > stale:
                    os.remove(lock_path)
                    continue
            except OSError:
                pass
            if time.monotonic() >= deadline:
                raise COFLockTimeout(
                    "Another COF is being generated right now. Please try again in a moment.")
            time.sleep(poll)
    try:
        yield
    finally:
        try:
            os.close(fd)
            os.remove(lock_path)
        except OSError:
            pass


# ── Orchestrator ─────────────────────────────────────────────────────
def generate_cof(d, workbook_path):
    """
    Validated data `d` + path to the stored active workbook.
    Atomically assigns the next number, mutates the stored workbook in place,
    and builds the Word doc. Returns a result dict (docx in memory).
    """
    if not os.path.exists(str(settings.COF_LETTERHEAD_PATH)):
        raise AssetMissing(
            "COF letterhead template is missing on the server "
            "(core/cof_assets/COF_LetterHead.docx).")

    with workbook_lock():
        info = get_next_cof_info(workbook_path)
        d = {**d,
             "cof_number": info["cof_number"],
             "claim_number": info["claim_number"],
             "optlog": info["optlog"]}
        try:
            wb = load_workbook(workbook_path)
        except PermissionError:
            raise WorkbookInUse(
                "The tracking workbook is open in Excel on the server. Close it and try again.")
        sheet_name = create_cof_sheet(wb, d)
        append_data_sheet(wb, d, info["serial"])
        try:
            wb.save(workbook_path)
        except PermissionError:
            raise WorkbookInUse(
                "Couldn't save — the tracking workbook is open in Excel on the server.")

    docx = build_word_doc(d)
    safe = re.sub(r'[^\w\s\-]', '', d["consignee_name"])[:40].strip()
    return {
        "cof_number": d["cof_number"],
        "claim_number": d["claim_number"],
        "optlog": d["optlog"],
        "serial": info["serial"],
        "sheet_name": sheet_name,
        "docx": docx,
        "docx_name": f"{d['cof_number']}_{safe}.docx",
    }


# ── History (reads the stored Data Sheet) ────────────────────────────
def load_history(workbook_path):
    headers = ["#", "LR Number", "Pickup Date", "Invoice Number", "Remark",
               "Dealer", "State", "Claim Amount", "CN Amount", "Delhivery Doc",
               "Doc Date", "Status Delhivery", "Remarks", "COF Date",
               "Status Optlog", "Ref Delhivery"]
    rows = []
    try:
        wb = load_workbook(workbook_path, read_only=True)
        ws = wb[DATA_SHEET]
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and isinstance(row[0], (int, float)):
                d = {}
                for i, h in enumerate(headers):
                    v = row[i] if i < len(row) else ""
                    if isinstance(v, datetime.datetime):
                        v = v.strftime("%d-%b-%Y")
                    d[h] = str(v) if v is not None else ""
                rows.append(d)
        wb.close()
    except Exception:
        return []
    return rows
